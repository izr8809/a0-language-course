from docx import Document
from docx.shared import Pt, Inches
import os
import re

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

characters = [
    ("male", "domin", "도민 (Domin)"),
    ("male", "hangyeol", "한결 (Hangyeol)"),
    ("male", "hyunjun", "현준 (Hyunjun)"),
    ("male", "siyun", "시윤 (Siyun) ♂"),
    ("male", "taeo", "태오 (Taeo)"),
    ("female", "minji", "민지 (Minji)"),
    ("female", "siyun", "서연 (Seoyeon)"),
    ("female", "yuna", "유나 (Yuna)"),
    ("female", "haneul", "하늘 (Haneul)"),
    ("female", "jia", "지아 (Jia)"),
]

themes = ["PI", "CR", "KD", "KP", "TR"]
base_path = "/Users/jaeseunglee/koko-script/a0/ko"

for gender, folder, display_name in characters:
    p = doc.add_paragraph()
    run = p.add_run(display_name)
    run.bold = True
    run.font.size = Pt(16)

    for theme in themes:
        file_path = os.path.join(base_path, gender, folder, f"{theme}.txt")
        if not os.path.exists(file_path):
            continue

        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        # Theme label
        p = doc.add_paragraph()
        run = p.add_run(f"[{theme}]")
        run.bold = True
        run.font.size = Pt(13)

        in_outro = False
        for line in lines:
            stripped = line.strip()

            if not stripped:
                in_outro = False
                continue

            # Skip outro lines
            if stripped.startswith('>'):
                in_outro = True
                continue
            if in_outro and not stripped.startswith('AI:') and not stripped.startswith('User:'):
                continue
            in_outro = False

            # Episode headers (# PI-1, # PI-2, etc.)
            if re.match(r'^#\s+(PI|CR|KD|KP|TR)-\d+', stripped) or re.match(r'^#\s+[^\s]', stripped):
                # Extract just the episode title
                title = stripped.lstrip('# ').strip()
                if title and not title.startswith('💖') and not title.startswith('✈') and not title.startswith('👔') and not title.startswith('🎬') and not title.startswith('🎤') and not title.startswith('☕'):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(8)
                    run = p.add_run(title)
                    run.bold = True
                    run.font.size = Pt(12)
                continue

            # Sim headers (### sim1, ## sim1, etc.)
            sim_match = re.match(r'^#{2,3}\s+sim\d+', stripped, re.IGNORECASE)
            if sim_match:
                sim_text = stripped.lstrip('# ').strip()
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                run = p.add_run(f"— {sim_text} —")
                run.bold = True
                run.font.size = Pt(11)
                continue

            # Skip all other metadata
            if stripped.startswith('#'):
                continue
            if stripped.startswith('**'):
                continue
            if stripped.startswith('- ') or stripped.startswith('|- '):
                continue
            if stripped.startswith('---'):
                continue
            if stripped.startswith('📍'):
                continue

            # AI and User lines only
            if stripped.startswith('AI:') or stripped.startswith('AI ('):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(11)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
            elif stripped.startswith('User:'):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.font.size = Pt(11)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)

    doc.add_page_break()

output_path = "/Users/jaeseunglee/koko-script/KOKO_A0_Lines_Only.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
