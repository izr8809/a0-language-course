from docx import Document
from docx.shared import Pt, Inches
import os

doc = Document()

# Remove default styles that might cause issues
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

characters = [
    ("male", "hangyeol", "한결 (Hangyeol)"),
    ("male", "hyunjun", "현준 (Hyunjun)"),
    ("male", "siyun", "시윤 (Siyun) ♂"),
    ("male", "taeo", "태오 (Taeo)"),
    ("male", "domin", "도민 (Domin)"),
    ("female", "siyun", "서연 (Seoyeon)"),
    ("female", "yuna", "유나 (Yuna)"),
    ("female", "haneul", "하늘 (Haneul)"),
    ("female", "jia", "지아 (Jia)"),
    ("female", "minji", "민지 (Minji)"),
]

themes = ["PI", "CR", "KD", "KP", "TR"]
base_path = "/Users/jaeseunglee/koko-script/a0/ko"

for gender, folder, display_name in characters:
    # Character name as big bold text
    p = doc.add_paragraph()
    run = p.add_run(f"━━━ {display_name} ━━━")
    run.bold = True
    run.font.size = Pt(16)

    for theme in themes:
        file_path = os.path.join(base_path, gender, folder, f"{theme}.txt")
        if not os.path.exists(file_path):
            continue

        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Theme label
        p = doc.add_paragraph()
        run = p.add_run(f"[{theme}]")
        run.bold = True
        run.font.size = Pt(13)

        # Just add the content as plain text with minimal formatting
        for line in content.split('\n'):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)

            stripped = line.strip()
            if not stripped:
                continue

            if stripped.startswith('AI:') or stripped.startswith('AI ('):
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(11)
            elif stripped.startswith('#'):
                run = p.add_run(stripped.lstrip('# '))
                run.bold = True
                run.font.size = Pt(12)
            elif stripped.startswith('**') and stripped.endswith('**'):
                run = p.add_run(stripped.strip('*'))
                run.bold = True
                run.font.size = Pt(11)
            elif stripped.startswith('>'):
                run = p.add_run(stripped)
                run.italic = True
                run.font.size = Pt(10)
            else:
                run = p.add_run(stripped)
                run.font.size = Pt(11)

    doc.add_page_break()

output_path = "/Users/jaeseunglee/koko-script/KOKO_A0_Scripts.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
